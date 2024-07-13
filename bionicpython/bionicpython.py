import argparse
import math
from docx import Document
import os
import argparse
import math
from docx import Document
import os
import spacy
import subprocess
import sys
import re




def process_word(word, ratio):
    # Calculate the number of characters to make bold
    num_chars = math.ceil(len(word) * ratio)
    
    # Split the word into the part that should be bold and the part that shouldn't
    bold_part = word[:num_chars]
    normal_part = word[num_chars:]
    
    # Create runs for the bold part and the normal part
    bold_run = (bold_part, True)
    normal_run = (normal_part, False)
    
    # Return the runs as a list
    return [bold_run, normal_run]




def process_document(doc_path, ratio):
    output_path = os.path.splitext(doc_path)[0] + '_modified.docx'

    # Check if the file path is for a .docx file
    if doc_path.endswith('.docx'):
        print("Already in docx format")
    elif doc_path.endswith('.pdf'):
        docx_path = doc_path.replace('.pdf', '.docx')
        try:
            # Run the conversion script as a subprocess
            subprocess.run([sys.executable, 'converter.py', '--pdf', doc_path, '--docx', docx_path])

            # Modify this command to run in the absolute path (with relation to this script) to converter.py
            subprocess.run([sys.executable, 'converter.py', '--pdf', doc_path, '--docx', docx_path])
            
            # Replace '.pdf' with '.docx' in the file path
            doc_path = docx_path
        except subprocess.CalledProcessError as e:
            print("Error converting PDF:", e)
            sys.exit(1)
    else:
        print("Files of this format are not supported yet")
        print("Please use either .pdf or .docx files")
        sys.exit()

    # Load the spacy model for word recognition (wrap in try-except)
    try:
        nlp = spacy.load('en_core_web_sm')
    except OSError as e:
        print("Error loading spaCy model:", e)
        sys.exit(1)

    try:
        # Open the .docx file
        word_doc = Document(doc_path)

        for paragraph in word_doc.paragraphs:
            for run in paragraph.runs:
                # Skip if the run is already bold
                if run.bold:
                    continue

                # Split the run text into words
                words = run.text.split(' ')
                words = [' ' + word if i != 0 else word for i, word in enumerate(words)]

                # Process each word
                new_runs = []
                for word in words:
                    # Use spacy to recognize the words
                    doc = nlp(word)
                    for token in doc:
                        # Bolden a ratio of the characters in the word
                        runs = process_word(token.text, ratio)
                        new_runs.extend(runs)

                # Clear the original run
                run.text = ''

                # Add new runs with the appropriate formatting
                for text, is_bold in new_runs:
                    new_run = paragraph.add_run(text)
                    new_run.bold = is_bold

        # Save the document (wrap in try-except)
        try:
            word_doc.save(output_path)
        except PermissionError as e:
            print("Error saving document:", e)
            sys.exit(1)

    except Exception as e:  # Catch any other unexpected errors
        print("Unexpected error processing document:", e)
        sys.exit(1)

    # Get the directory and filename from the input path
    dir_name, file_name = os.path.split(doc_path)

    # Prepend 'processed_' to the filename
    new_file_name = 'processed_' + file_name

    # Combine the directory and new filename to get the output path
    # output_path = os.path.splitext(doc_path)[0] + '_modified.docx'
    print(output_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('doc_path', help='Path to the .docx file')
    parser.add_argument('--ratio', type=float, default=0.5, help='Ratio of characters to make bold')
    args = parser.parse_args()

    process_document(args.doc_path, args.ratio)

if __name__ == '__main__':
    main()

# def process_document(doc_path, ratio):
#     # Check if the file path is for a .docx file
#     if not doc_path.endswith('.docx'):
#         print("Invalid file format. Please provide a .docx file.")
#         return
    
#     # Load the spacy model for word recognition
#     nlp = spacy.load('en_core_web_sm')
    
#     # Open the .docx file
#     doc = Document(doc_path)

#     for paragraph in doc.paragraphs:
#         for run in paragraph.runs:
#             # Skip if the run is already bold
#             if run.bold:
#                 continue

#             # Split the run text into words
#             words = run.text.split()

#             # Process each word
#             new_runs = []
#             for word in words:
#                 # Use spacy to recognize the words
#                 doc = nlp(word)
#                 for token in doc:
#                     # Bolden a ratio of the characters in the word
#                     runs = process_word(token.text, ratio)
#                     new_runs.extend(runs)

#             # Clear the original run
#             run.text = ''

#             # Add new runs with the appropriate formatting
#             for text, is_bold in new_runs:
#                 new_run = paragraph.add_run(text)
#                 new_run.bold = is_bold


#     # Get the directory and filename from the input path
#     dir_name, file_name = os.path.split(doc_path)

#     # Prepend 'processed_' to the filename
#     new_file_name = 'processed_' + file_name

#     # Combine the directory and new filename to get the output path
#     output_path = os.path.join(dir_name, new_file_name)

#     doc.save(output_path)


# def process_word(word, ratio):
#     # Calculate the number of characters to make bold
#     num_chars = math.ceil(len(word) * ratio)
    
#     # Make the specified number of characters bold
#     bold_word = word[:num_chars] + word[num_chars:].bold()
    
#     return bold_word

# def process_document(doc_path, ratio):
#     # Check if the file path is for a .docx file
#     if not doc_path.endswith('.docx'):
#         print("Invalid file format. Please provide a .docx file.")
#         return
    
#     # Load the spacy model for word recognition
#     nlp = spacy.load('en_core_web_sm')
    
#     # Open the .docx file
#     doc = Document(doc_path)

#     for paragraph in doc.paragraphs:
#         for run in paragraph.runs:
#             # Skip if the run is already bold
#             if run.bold:
#                 continue

#             # Split the run text into words
#             words = run.text.split()

#             # Process each word
#             new_words = []
#             for word in words:
#                 # Use spacy to recognize the words
#                 doc = nlp(word)
#                 for token in doc:
#                     # Bolden a ratio of the characters in the word
#                     bold_word = process_word(token.text, ratio)
#                     new_words.append(bold_word)

#             # Replace the run text with the processed words
#             run.text = ' '.join(new_words)
#             print("processed another paragraph...")

#     # Get the directory and filename from the input path
#     dir_name, file_name = os.path.split(doc_path)

#     # Prepend 'processed_' to the filename
#     new_file_name = 'processed_' + file_name

#     # Combine the directory and new filename to get the output path
#     output_path = os.path.join(dir_name, new_file_name)

#     doc.save(output_path)


# def main():
#     parser = argparse.ArgumentParser()
#     parser.add_argument('doc_path', help='Path to the .docx file')
#     parser.add_argument('--ratio', type=float, default=0.5, help='Ratio of characters to make bold')
#     args = parser.parse_args()

#     process_document(args.doc_path, args.ratio)

# if __name__ == '__main__':
#     main()




# python bionic.py /path/to/your/document.docx 























# def process_word(word, ratio):
#     # Calculate the number of characters to make bold
#     num_bold_chars = math.ceil(len(word) * ratio)

#     # Create a new run for each part of the word
#     bold_run = word[:num_bold_chars]
#     normal_run = word[num_bold_chars:]

#     return bold_run, normal_run

# # def process_document(doc_path, ratio):
# #     doc = Document(doc_path)

# #     for paragraph in doc.paragraphs:
# #         for run in paragraph.runs:
# #             # Skip if the run is already bold
# #             if run.bold:
# #                 continue

# #             # Split the run text into words
# #             words = run.text.split()

# #             # Process each word
# #             new_words = []
# #             for word in words:
# #                 bold_run, normal_run = process_word(word, ratio)
# #                 new_words.append(bold_run)
# #                 new_words.append(normal_run)

# #             # Replace the run text with the processed words
# #             run.text = ' '.join(new_words)

# #     doc.save('processed_' + doc_path)

# # def process_document(doc_path, ratio):
# #     doc = Document(doc_path)

# #     for paragraph in doc.paragraphs:
# #         for run in paragraph.runs:
# #             # Skip if the run is already bold
# #             if run.bold:
# #                 continue

# #             # Split the run text into words
# #             words = run.text.split()

# #             # Process each word
# #             new_words = []
# #             for word in words:
# #                 bold_run, normal_run = process_word(word, ratio)
# #                 new_words.append(bold_run)
# #                 new_words.append(normal_run)

# #             # Replace the run text with the processed words
# #             run.text = ' '.join(new_words)

# #     # Get the directory and filename from the input path
# #     dir_name, file_name = os.path.split(doc_path)

# #     # Prepend 'processed_' to the filename
# #     new_file_name = 'processed_' + file_name

# #     # Combine the directory and new filename to get the output path
# #     output_path = os.path.join(dir_name, new_file_name)

# #     doc.save(output_path)

# def process_document(doc_path, ratio):
#     doc = Document(doc_path)
#     new_doc = Document()

#     for paragraph in doc.paragraphs:
#         new_paragraph = new_doc.add_paragraph()

#         for run in paragraph.runs:
#             # Skip if the run is already bold
#             if run.bold:
#                 continue

#             # Split the run text into words
#             words = run.text.split()

#             # Process each word
#             for word in words:
#                 bold_run_text, normal_run_text = process_word(word, ratio)

#                 # Add new runs to the new paragraph
#                 if bold_run_text:
#                     bold_run = new_paragraph.add_run(bold_run_text)
#                     bold_run.bold = True

#                 if normal_run_text:
#                     normal_run = new_paragraph.add_run(normal_run_text)

#     # Get the directory and filename from the input path
#     dir_name, file_name = os.path.split(doc_path)

#     # Prepend 'processed_' to the filename
#     new_file_name = 'processed_' + file_name

#     # Combine the directory and new filename to get the output path
#     output_path = os.path.join(dir_name, new_file_name)

#     new_doc.save(output_path)


# def main():
#     parser = argparse.ArgumentParser()
#     parser.add_argument('doc_path', help='Path to the .docx file')
#     parser.add_argument('--ratio', type=float, default=0.5, help='Ratio of characters to make bold')
#     args = parser.parse_args()

#     process_document(args.doc_path, args.ratio)

# if __name__ == '__main__':
#     main()




